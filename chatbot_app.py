import streamlit as st
import openai
import os
import json
import hashlib
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
from docx import Document
import glob
import io

# Configuração da chave da API
openai.api_key = os.getenv("OPENAI_API_KEY")

# Funções de gerenciamento de usuários e autenticação
def load_users():
    return json.load(open('users.json', 'r')) if os.path.exists('users.json') else {}

def save_users(users):
    json.dump(users, open('users.json', 'w'))

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def authenticate(username, password):
    users = load_users()
    return username in users and users[username] == hash_password(password)

def register_user(username, password):
    users = load_users()
    if username in users:
        return False
    users[username] = hash_password(password)
    save_users(users)
    return True

# Funções para extrair texto de diferentes tipos de arquivo
def extract_text_from_docx(file):
    if isinstance(file, str):
        doc = Document(file)
    else:
        doc = Document(file)
    return " ".join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_txt(file):
    if isinstance(file, str):
        with open(file, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        return file.getvalue().decode("utf-8")

def extract_text_from_pdf(file):
    if isinstance(file, str):
        with open(file, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            return " ".join([page.extract_text() for page in pdf_reader.pages])
    else:
        pdf_reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() for page in pdf_reader.pages])

def extract_text(file):
    try:
        if isinstance(file, str):
            file_extension = os.path.splitext(file)[1].lower()
        else:
            file_extension = os.path.splitext(file.name)[1].lower()
        
        if file_extension == '.docx':
            return extract_text_from_docx(file)
        elif file_extension == '.txt':
            return extract_text_from_txt(file)
        elif file_extension == '.pdf':
            return extract_text_from_pdf(file)
        else:
            return "Formato de arquivo não suportado."
    except Exception as e:
        st.warning(f"Erro ao ler o arquivo: {e}")
        return ""

# Função para carregar a base de conhecimento
@st.cache_data(show_spinner=False)
def load_knowledge_base():
    knowledge_base = {}
    for ext in ['docx', 'txt', 'pdf']:
        for file_path in glob.glob(f"knowledge_base/*.{ext}"):
            content = extract_text(file_path)
            if content:
                knowledge_base[os.path.basename(file_path)] = content
    return knowledge_base

# Função para encontrar as seções mais relevantes
def get_most_relevant_sections(query, knowledge_base, top_n=3):
    if not knowledge_base:
        return "Base de conhecimento vazia ou não carregada corretamente."
    
    documents = list(knowledge_base.values())
    doc_names = list(knowledge_base.keys())
    
    vectorizer = TfidfVectorizer().fit_transform([query] + documents)
    similarities = cosine_similarity(vectorizer[0], vectorizer[1:]).flatten()
    
    relevant_indices = similarities.argsort()[-top_n:][::-1]
    return "\n\n".join([f"{doc_names[i]}: {documents[i]}" for i in relevant_indices])

# Função para buscar informações na base de conhecimento
def search_knowledge_base(query):
    knowledge_base = load_knowledge_base()
    relevant_content = get_most_relevant_sections(query, knowledge_base)
    
    if relevant_content == "Base de conhecimento vazia ou não carregada corretamente.":
        return relevant_content
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Você é um assistente especializado em estética automotiva. Seu objetivo é fornecer informações precisas e úteis sobre como montar, gerir, construir e gerenciar uma estética automotiva. Responda de forma didática e humanizada, baseando-se no conteúdo fornecido e em seu conhecimento especializado."},
                {"role": "user", "content": f"Com base no seguinte conteúdo e em seu conhecimento sobre estética automotiva, responda à pergunta de forma precisa e humanizada: '{query}'\n\nConteúdo: {relevant_content[:4000]}"}
            ]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        return f"Erro ao se comunicar com a API OpenAI: {e}"

# Estilo CSS personalizado
def local_css():
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');
        
        body {
            color: white;
            background-color: #8B0000; /* Vermelho Rubi */
            font-family: 'Roboto', sans-serif;
        }
        .main {
            padding: 2rem;
            max-width: 800px;
            margin: 0 auto;
            background-color: rgba(255, 255, 255, 0.1);
            border-radius: 15px;
            box-shadow: 0 10px 20px rgba(0, 0, 0, 0.2);
        }
        .stTextInput>div>div>input {
            background-color: white;
            color: #333;
            border: 1px solid #8B0000;
            border-radius: 5px;
            padding: 10px;
        }
        .stButton>button {
            width: 100%;
            padding: 0.7rem 1rem;
            background-color: white;
            color: #8B0000;
            border: 2px solid #8B0000;
            border-radius: 5px;
            font-weight: bold;
            transition: all 0.3s ease;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        .stButton>button:hover {
            background-color: #8B0000;
            color: white;
            box-shadow: 0 6px 8px rgba(0, 0, 0, 0.2);
        }
        .chat-message {
            padding: 1rem;
            border-radius: 10px;
            margin-bottom: 1rem;
            display: flex;
            flex-direction: row;
            align-items: flex-start;
            flex-wrap: wrap;
            background-color: rgba(255, 255, 255, 0.1);
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        }
        .chat-message.user {
            background-color: rgba(255, 255, 255, 0.2);
        }
        .chat-message.bot {
            background-color: rgba(0, 0, 0, 0.3);
        }
        .chat-message .avatar {
            width: 50px;
            height: 50px;
            margin-right: 1rem;
        }
        .chat-message .avatar img {
            width: 100%;
            height: 100%;
            border-radius: 50%;
            object-fit: cover;
            border: 2px solid white;
        }
        .chat-message .message {
            flex: 1;
            padding: 0;
            color: white;
            word-wrap: break-word;
        }
        .stTextInput>div>div>input::placeholder {
            color: #888;
        }
        h1, h2, h3 {
            color: white;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.5);
        }
        .subtitle {
            font-size: 1.2rem;
            color: #FFD700;
            margin-bottom: 2rem;
        }
        .login-container {
            background-color: rgba(0, 0, 0, 0.6);
            padding: 2rem;
            border-radius: 15px;
            box-shadow: 0 8px 16px rgba(0, 0, 0, 0.3);
        }
        /* Responsividade */
        @media (max-width: 768px) {
            .chat-message {
                flex-direction: column;
            }
            .chat-message .avatar {
                width: 40px;
                height: 40px;
                margin-bottom: 0.5rem;
            }
            .chat-message .message {
                width: 100%;
            }
        }
        /* Ajustes para telas menores */
        @media (max-width: 480px) {
            .stButton>button {
                padding: 0.5rem 0.7rem;
                font-size: 0.9rem;
            }
            .chat-message {
                padding: 0.7rem;
            }
        }
    </style>
    """, unsafe_allow_html=True)

# Interface do Streamlit
def main():
    st.set_page_config(page_title="Assistente de IA LG", layout="wide")
    local_css()

    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False

    if not st.session_state.logged_in:
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            st.markdown("<div class='login-container'>", unsafe_allow_html=True)
            st.title("Assistente de IA LG 🚗")
            st.markdown("<p class='subtitle'>Sua Solução Inteligente em Estética Automotiva</p>", unsafe_allow_html=True)
            username = st.text_input("Usuário")
            password = st.text_input("Senha", type="password")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Entrar", key="login"):
                    if authenticate(username, password):
                        st.session_state.logged_in = True
                        st.rerun()
                    else:
                        st.error("Usuário ou senha incorretos.")
            with col2:
                if st.button("Cadastrar", key="register"):
                    if register_user(username, password):
                        st.success("Usuário registrado com sucesso!")
                    else:
                        st.error("Usuário já existe.")
            st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.title("Assistente de IA LG - Especialista em Estética Automotiva 🚗")
        
        if 'messages' not in st.session_state:
            st.session_state.messages = []
        
        if 'knowledge_base' not in st.session_state:
            st.session_state.knowledge_base = load_knowledge_base()

        uploaded_file = st.file_uploader("Faça upload de um novo arquivo para a base de conhecimento", type=['txt', 'pdf', 'docx'])
        if uploaded_file is not None:
            content = extract_text(uploaded_file)
            if content:
                st.session_state.knowledge_base[uploaded_file.name] = content
                st.success(f"Arquivo {uploaded_file.name} adicionado à base de conhecimento.")

        for message in st.session_state.messages:
            with st.container():
                st.markdown(f"""
                <div class="chat-message {message['role']}">
                    <div class="avatar">
                        <img src="https://i.imgur.com/{'Rnj9Ry1.png' if message['role'] == 'user' else 'mFQOKAq.png'}">
                    </div>
                    <div class="message">{message['content']}</div>
                </div>
                """, unsafe_allow_html=True)

        prompt = st.text_input("Digite sua pergunta sobre estética automotiva aqui...", key="user_input")
        col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
        with col2:
            if st.button("Enviar", key="send"):
                if prompt:
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    response = search_knowledge_base(prompt)
                    st.session_state.messages.append({"role": "bot", "content": response})
                    st.rerun()
        with col3:
            if st.button("Limpar", key="clear"):
                st.session_state.messages = []
                st.rerun()
        with col4:
            if st.button("Logout", key="logout"):
                st.session_state.logged_in = False
                st.rerun()

if __name__ == "__main__":
    main()