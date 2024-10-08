import streamlit as st
import openai
import os
import json
import hashlib
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
from docx import Document
import glob
import io

# Configuração da chave da API
openai.api_key = os.getenv("OPENAI_API_KEY")

# Funções de gerenciamento de usuários e autenticação
def load_users():
    return json.load(open('users.json', 'r')) if os.path.exists('users.json') else {}

def save_users(users):
    json.dump(users, open('users.json', 'w'))

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def authenticate(username, password):
    users = load_users()
    return username in users and users[username] == hash_password(password)

def register_user(username, password):
    users = load_users()
    if username in users:
        return False
    users[username] = hash_password(password)
    save_users(users)
    return True

# Funções para extrair texto de diferentes tipos de arquivo
def extract_text_from_docx(file):
    if isinstance(file, str):
        doc = Document(file)
    else:
        doc = Document(file)
    return " ".join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_txt(file):
    if isinstance(file, str):
        with open(file, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        return file.getvalue().decode("utf-8")

def extract_text_from_pdf(file):
    if isinstance(file, str):
        with open(file, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            return " ".join([page.extract_text() for page in pdf_reader.pages])
    else:
        pdf_reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() for page in pdf_reader.pages])

def extract_text(file):
    try:
        if isinstance(file, str):
            file_extension = os.path.splitext(file)[1].lower()
        else:
            file_extension = os.path.splitext(file.name)[1].lower()
        
        if file_extension == '.docx':
            return extract_text_from_docx(file)
        elif file_extension == '.txt':
            return extract_text_from_txt(file)
        elif file_extension == '.pdf':
            return extract_text_from_pdf(file)
        else:
            return "Formato de arquivo não suportado."
    except Exception as e:
        st.warning(f"Erro ao ler o arquivo: {e}")
        return ""

# Função para carregar a base de conhecimento
@st.cache_data(show_spinner=False)
def load_knowledge_base():
    knowledge_base = {}
    for ext in ['docx', 'txt', 'pdf']:
        for file_path in glob.glob(f"knowledge_base/*.{ext}"):
            content = extract_text(file_path)
            if content:
                knowledge_base[os.path.basename(file_path)] = content
    return knowledge_base

# Função para encontrar as seções mais relevantes
def get_most_relevant_sections(query, knowledge_base, top_n=3):
    if not knowledge_base:
        return "Base de conhecimento vazia ou não carregada corretamente."
    
    documents = list(knowledge_base.values())
    doc_names = list(knowledge_base.keys())
    
    vectorizer = TfidfVectorizer().fit_transform([query] + documents)
    similarities = cosine_similarity(vectorizer[0], vectorizer[1:]).flatten()
    
    relevant_indices = similarities.argsort()[-top_n:][::-1]
    return "\n\n".join([f"{doc_names[i]}: {documents[i]}" for i in relevant_indices])

# Função para buscar informações na base de conhecimento
def search_knowledge_base(query):
    knowledge_base = load_knowledge_base()
    relevant_content = get_most_relevant_sections(query, knowledge_base)
    
    if relevant_content == "Base de conhecimento vazia ou não carregada corretamente.":
        return relevant_content
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Você é um assistente especializado em estética automotiva. Seu objetivo é fornecer informações precisas e úteis sobre como montar, gerir, construir e gerenciar uma estética automotiva. Responda de forma didática e humanizada, baseando-se no conteúdo fornecido e em seu conhecimento especializado."},
                {"role": "user", "content": f"Com base no seguinte conteúdo e em seu conhecimento sobre estética automotiva, responda à pergunta de forma precisa e humanizada: '{query}'\n\nConteúdo: {relevant_content[:4000]}"}
            ]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        return f"Erro ao se comunicar com a API OpenAI: {e}"

# Estilo CSS personalizado
def local_css():
    st.markdown("""
    <style>
        # Seu código CSS personalizado aqui
    </style>
    """, unsafe_allow_html=True)

# Interface do Streamlit
def main():
    st.set_page_config(page_title="Assistente de IA LG", layout="wide")
    local_css()

    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False

    if not st.session_state.logged_in:
        # Código de login aqui
        pass
    else:
        # Código da interface do aplicativo aqui
        pass

if __name__ == "__main__":
    main()
